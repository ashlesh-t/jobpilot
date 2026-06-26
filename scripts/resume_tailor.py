"""Resume tailor — DOCX fallback only.

Used when the user has no LaTeX source. Builds a keyword-optimised DOCX from profile.json
and the job JD. Primary tailoring (LaTeX editing + tectonic compile) is handled directly by
Claude in the job-tailor skill.

Guards:
  * self-caps at top_n_tailor (default 5) per calendar run via /tmp counter
  * skips if this job_id already has a tailored_resume_path in SQLite
"""
from __future__ import annotations

import json
import os
import re
import sqlite3
import sys
from datetime import datetime, timezone
from pathlib import Path

FILTERED_IN = "/tmp/jobpilot_filtered.json"
RUN_COUNTER = "/tmp/jobpilot_tailor_count.txt"


def jobpilot_dir() -> Path:
    raw = os.environ.get("JOBPILOT_DIR", "~/.claude/job-hunt-ai")
    return Path(os.path.expanduser(raw))


def db_path() -> Path:
    return jobpilot_dir() / "cache" / "jobs.sqlite"


def load_json(path, default):
    try:
        return json.loads(Path(path).read_text())
    except Exception:
        return default


def find_job(job_id: str) -> dict | None:
    for job in load_json(FILTERED_IN, []):
        if job.get("job_id") == job_id:
            return job
    return None


def already_tailored(job_id: str) -> bool:
    if not db_path().exists():
        return False
    try:
        conn = sqlite3.connect(str(db_path()))
        row = conn.execute(
            "SELECT tailored_resume_path FROM jobs_seen WHERE job_id=?", (job_id,)
        ).fetchone()
        conn.close()
        return bool(row and row[0])
    except sqlite3.Error:
        return False


def run_count() -> int:
    try:
        return int(Path(RUN_COUNTER).read_text().strip())
    except Exception:
        return 0


def bump_count() -> None:
    Path(RUN_COUNTER).write_text(str(run_count() + 1))


def set_tailored_path(job_id: str, path: str) -> None:
    if not db_path().exists():
        return
    try:
        conn = sqlite3.connect(str(db_path()))
        now = datetime.now(timezone.utc).isoformat()
        conn.execute(
            "INSERT INTO jobs_seen (job_id, tailored_resume_path, first_seen, last_seen, status) "
            "VALUES (?,?,?,?, 'active') "
            "ON CONFLICT(job_id) DO UPDATE SET tailored_resume_path=excluded.tailored_resume_path,"
            " last_seen=excluded.last_seen",
            (job_id, path, now, now),
        )
        conn.commit()
        conn.close()
    except sqlite3.Error as exc:
        print(f"[tailor] db update failed: {exc}", file=sys.stderr)


def safe_company(job: dict) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "-", (job.get("company") or "company")).strip("-").lower()


def build_docx(job: dict, profile: dict, matched_skills: list, out_path: Path) -> None:
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    name = profile.get("name") or "Your Name"
    doc.add_heading(name, level=0)
    contact = profile.get("email", "")
    if contact:
        doc.add_paragraph(contact)

    doc.add_heading("Objective", level=1)
    doc.add_paragraph(
        f"{profile.get('experience_years', 0)}+ yr engineer targeting the "
        f"{job.get('role', 'Software Engineer')} role at {job.get('company', '')}, "
        f"bringing hands-on strengths in the technologies this team uses."
    )

    # Lead with JD-matched skills
    all_skills = matched_skills + [s for s in profile.get("skills", []) if s not in matched_skills]
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(all_skills) if all_skills else "See base resume")

    if profile.get("roles_held"):
        doc.add_heading("Experience", level=1)
        for r in profile["roles_held"]:
            doc.add_paragraph(str(r), style="List Bullet")

    if profile.get("projects"):
        doc.add_heading("Projects", level=1)
        for p in profile["projects"]:
            text = p if isinstance(p, str) else json.dumps(p)
            doc.add_paragraph(text, style="List Bullet")

    edu = profile.get("education", {})
    if edu:
        doc.add_heading("Education", level=1)
        doc.add_paragraph(", ".join(str(v) for v in edu.values()))

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = run.font.size or Pt(11)

    doc.save(str(out_path))


def tailor(job_id: str, matched_skills: list | None = None) -> str | None:
    prefs = load_json(jobpilot_dir() / "options" / "preferences.json", {})
    profile = load_json(jobpilot_dir() / "cache" / "profile.json", {})
    cap = int(prefs.get("top_n_tailor", 5))

    job = find_job(job_id)
    if not job:
        print(f"[tailor] job {job_id} not found in {FILTERED_IN}", file=sys.stderr)
        return None
    if already_tailored(job_id):
        print(f"[tailor] {job_id} already tailored — skipping.")
        return None
    if run_count() >= cap:
        print(f"[tailor] per-run cap of {cap} reached — skipping {job_id}.")
        return None

    tailored_dir = jobpilot_dir() / "resumes" / "tailored"
    tailored_dir.mkdir(parents=True, exist_ok=True)
    stem = f"{safe_company(job)}-{job_id}"
    out = tailored_dir / f"{stem}.docx"

    build_docx(job, profile, matched_skills or [], out)
    set_tailored_path(job_id, str(out))
    bump_count()
    print(f"[tailor] wrote {out}")
    return str(out)


def main() -> None:
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    if not args:
        print("usage: python3 resume_tailor.py <job_id> [--matched-skills skill1,skill2]",
              file=sys.stderr)
        sys.exit(1)

    matched_skills = []
    if "--matched-skills" in sys.argv:
        idx = sys.argv.index("--matched-skills")
        if idx + 1 < len(sys.argv):
            matched_skills = [s.strip() for s in sys.argv[idx + 1].split(",") if s.strip()]

    tailor(args[0], matched_skills)


if __name__ == "__main__":
    main()
